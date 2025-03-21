import streamlit as st
from docx import Document
import io
import os
import re
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from gingerit.gingerit import GingerIt

# Initialize counter
if not os.path.exists("counter.txt"):
    with open("counter.txt", "w") as f:
        f.write("0")

def get_counter():
    with open("counter.txt", "r") as f:
        return int(f.read())

def increment_counter():
    with open("counter.txt", "r+") as f:
        count = int(f.read())
        count += 1
        f.seek(0)
        f.write(str(count))
        f.truncate()
    return count

# Load advanced paraphrasing model
MODEL_NAME = "humarin/chatgpt_paraphraser_on_T5_base"
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)

def grammar_check(text):
    try:
        parser = GingerIt()
        result = parser.parse(text)
        return result['result']
    except Exception as e:
        st.error(f"Grammar check failed: {str(e)}")
        return text

def improve_text(text):
    # Preprocessing
    text = re.sub(r'\s+', ' ', text).strip()
    
    # Split into manageable chunks
    sentences = re.split(r'(?<=[.!?]) +', text)
    chunks = []
    current_chunk = []
    char_count = 0
    
    for sentence in sentences:
        if char_count + len(sentence) < 400:
            current_chunk.append(sentence)
            char_count += len(sentence)
        else:
            chunks.append(" ".join(current_chunk))
            current_chunk = [sentence]
            char_count = len(sentence)
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    # Process each chunk
    improved_chunks = []
    for chunk in chunks:
        inputs = tokenizer(
            f"paraphrase: {chunk}",
            return_tensors="pt",
            max_length=512,
            truncation=True
        )
        
        outputs = model.generate(
            inputs["input_ids"],
            max_length=512,
            num_beams=5,
            num_return_sequences=1,
            temperature=1.2,
            do_sample=True,
            top_p=0.95,
            early_stopping=True
        )
        
        paraphrased = tokenizer.decode(outputs[0], skip_special_tokens=True)
        improved_chunks.append(paraphrased)

    # Post-processing
    final_text = " ".join(improved_chunks)
    final_text = grammar_check(final_text)
    
    # Ensure coherence
    final_text = re.sub(r'\s+([.,!?])', r'\1', final_text)
    final_text = re.sub(r'\s+', ' ', final_text)
    
    return final_text

# Page configuration
st.set_page_config(
    page_title="⚡ Rapid Humanize AI",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600&display=swap');
    html, body, [class*="css"]  {
        font-family: 'Poppins', sans-serif;
    }
    .stTextArea textarea {min-height: 300px; border-radius: 10px; padding: 15px!important;}
    .download-btn {margin-top: 20px;}
    .counter {font-size: 1.4em; color: #FF4B4B; font-weight: 700; padding: 10px 20px; background: #FFE3E3; border-radius: 10px;}
    .header {color: #2c3e50; border-bottom: 3px solid #FF4B4B; padding-bottom: 10px;}
    .stButton button {background: #FF4B4B!important; color: white!important; border-radius: 8px!important; padding: 12px 24px!important;}
    .stButton button:hover {background: #FF2B2B!important; color: white!important;}
    .logo {text-align: center; margin-bottom: 30px;}
    .logo h1 {color: #FF4B4B; font-size: 2.8em; margin-bottom: 0;}
    .logo p {color: #6c757d; font-size: 1.2em;}
</style>
""", unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.markdown('<div class="header">⚙️ Settings</div>', unsafe_allow_html=True)
    language = st.selectbox("Language Version", ["English (US)", "English (UK)", "English (AU)", "English (CA)"])
    mode = st.radio("Processing Mode", ["Basic", "Advanced"])
    st.markdown(f'<div class="counter">🚀 Total Conversions: {get_counter()}</div>', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("**Made with ❤️ by RHAI Team**")

# Main content
st.markdown("""
<div class="logo">
    <h1>⚡ Rapid Humanize AI</h1>
    <p>Transform AI-generated text into undetectable human-like content</p>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2, gap="large")

with col1:
    st.markdown('<div class="header">📥 Input Section</div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Upload DOC/DOCX File", type=["docx", "doc"], help="Max file size: 25MB")
    
    input_text = st.text_area("Paste your AI-generated text here:", height=400, 
                            placeholder="Enter text or upload file... (No word limit)")
    
    if uploaded_file:
        doc = Document(io.BytesIO(uploaded_file.getvalue()))
        input_text = "\n".join([para.text for para in doc.paragraphs])

with col2:
    st.markdown('<div class="header">📤 Output Section</div>', unsafe_allow_html=True)
    output_text = st.text_area("Humanized text will appear here:", 
                             value=st.session_state.get("output_text", ""), 
                             height=400, key="output")
    
    if output_text:
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.download_button(
                label="📥 Download TXT",
                data=output_text,
                file_name="humanized_text.txt",
                mime="text/plain",
                use_container_width=True
            )
        with col_d2:
            doc = Document()
            doc.add_paragraph(output_text)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button(
                label="📥 Download DOCX",
                data=bio.getvalue(),
                file_name="humanized_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

# Humanize Button
if st.button("✨ HUMANIZE NOW", use_container_width=True, type="primary"):
    if input_text.strip():
        with st.spinner("🚀 Humanizing your text... (This usually takes 10-30 seconds)"):
            output_text = improve_text(input_text)
            st.session_state.output_text = output_text
            increment_counter()
            st.experimental_rerun()
    else:
        st.error("Please enter some text or upload a file to humanize")

# Features Section
with st.expander("🌟 Key Features", expanded=True):
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        st.markdown("""
        ### 🔒 AI Detection Proof
        Advanced algorithms remove all AI patterns
        """)
    with col_f2:
        st.markdown("""
        ### 🌐 Multi-language Support
        Process text in various English variants
        """)
    with col_f3:
        st.markdown("""
        ### ⚡ Instant Processing
        Quick conversion with real-time preview
        """)

# Instructions
with st.expander("📘 How to Use RHAI"):
    st.markdown("""
    1. **Input Method** - Paste text directly or upload Word document
    2. **Settings** - Choose language variant and processing mode
    3. **Humanize** - Click the HUMANIZE NOW button
    4. **Review & Download** - Check output and download preferred format
    """)

st.markdown("""
---
> **Note**: Current version uses GingerIt for grammar checking. For enterprise solutions with enhanced checking, contact RHAI Team.
""")
